# -*- coding: utf-8 -*-
"""在产品说明书中补充 2.2.2～2.2.4（数据概览子页面），并标注 2.2.1。"""
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "\u7269\u4e1a\u6210\u672c\u7ba1\u7406\u7cfb\u7edf\u4ea7\u54c1\u8bf4\u660e\u4e66.docx"
OUT_PATH = ROOT / "\u7269\u4e1a\u6210\u672c\u7ba1\u7406\u7cfb\u7edf\u4ea7\u54c1\u8bf4\u660e\u4e66_\u66f4\u65b0.docx"


def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def main():
    doc = Document(str(DOC_PATH))

    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "物业成本汇总页面":
            p.text = "2.2.1 物业成本汇总页面"
            break

    insert_after_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "表格定义":
            insert_after_idx = i
            break
    if insert_after_idx is None:
        raise SystemExit("未找到「表格定义」段落，未写入。")

    anchor = doc.paragraphs[insert_after_idx]

    lines = [
        "2.2.2 成本还原页面",
        "",
        "页面展示",
        "",
        "数据汇总区",
        "页面顶部数据汇总区展示 3 项核心统计指标，统计口径为成本还原表中各一级科目在「内控值」行上的金额合计，包括：全年合计（内控值汇总）、已发生&支付（内控值汇总）、未发生&未支付（内控值汇总）。金额保留 2 位小数，以卡片形式展示。",
        "",
        "筛选区",
        "页面上方为查询条件区，布局与物业成本汇总页面一致：支持选择项目、统计周期（日期范围）。后续版本可对接后端接口后按筛选条件刷新表格数据；当前前端用于版式与数据展示样例。",
        "",
        "成本还原表区域",
        "筛选区下方为成本还原表展示区域，采用二级表头、横向分组的宽表结构。",
        "行结构：按一级成本科目展开；每个科目占三行，对应三种口径——内控值、过程管控值、实际&预计值。一级科目名称在「内控值」行以加粗显示。",
        "列结构：左侧固定列为「一级科目」「口径」；其后为全年维度三列——全年合计、已发生&支付、未发生&未支付；再按自然月（YYYY-MM）分组，每月下列三列——已支出金额、预计未支付金额、未支付金额。",
        "展示规则：金额类字段右对齐，统一保留 2 位小数；源数据为空或错误时单元格显示「—」；表格横向滚动浏览，整表不分页展示。",
        "",
        "界面截图说明",
        "系统实现页面路径：数据概览 → 成本还原。截图版式应与「页标题 + 顶部指标卡片 + 筛选卡片 + 底部表格卡片」分区一致，可参考前端实现文件 client/src/pages/overview/CostRestore.jsx。",
        "",
        "2.2.3 明细台账页面",
        "",
        "页面展示",
        "",
        "数据汇总区",
        "页面顶部展示表头辅助信息与 5 项年度汇总指标：项目名称、不含税合计（表头）取自台账表头；五项指标取自台账「合计」行在全年区块下的数值，包括：预估年度金额（合计）、内控金额（合计）、预估发生金额（合计）、实际发生金额（合计）、实付金额（合计）。金额保留 2 位小数。",
        "",
        "筛选区",
        "与成本汇总页风格一致，提供筛选条件卡片：选择项目、选择时间（起止日期）。用于后续与接口联动刷新明细数据。",
        "",
        "明细台账表区域",
        "筛选区下方为明细台账表。左侧固定三列：成本费用科目、不含税金额、税率（0～1 的小数按百分比展示，如 0.06 显示为 6.00%）。",
        "中间为「全年」分组（表头名称与导入模板一致，如「2023全年」），下列五列：预估年度金额、内控金额、预估发生金额、实际发生金额、实付金额。",
        "其后按自然月（如 2022-12～2023-12）分组，每月下列五列：内控金额、预估发生金额、实际发生金额、实付金额、未支付金额。",
        "展示规则：金额右对齐、保留 2 位小数；无数据显示「—」；「合计」行加粗并采用与系统一致的合计行底色以便识别；表格支持横向滚动。",
        "",
        "界面截图说明",
        "系统实现页面路径：数据概览 → 明细台账。版式参考 client/src/pages/overview/DetailLedger.jsx；数据可由明细台账 Excel 经脚本导出为 detailLedgerData.json 维护。",
        "",
        "2.2.4 能耗台账页面",
        "",
        "页面展示",
        "",
        "数据汇总区",
        "页面顶部展示表头名称（能耗明细表）及 6 项核心指标，取自首条汇总数据行在「全年」相关区块中的数值：全年（已发生+预估发生金额）下的预算值、管理值、完成率；全年滚动已发生下的预算值、实际发生金额、剩余差额。完成率以小数口径存储，界面以百分比展示（保留 2 位小数）。",
        "",
        "筛选区",
        "提供筛选条件卡片：选择项目、选择时间（起止日期），布局与上述概览子页面一致。",
        "",
        "能耗明细表区域",
        "筛选区下方为能耗明细表。左侧固定三列：指标名称、能耗类型、指标类型。",
        "其后为三个「全年」分组表头（与 Excel 一致）：全年（已发生+预估发生金额）——预算值、管理值、完成率；全年滚动已发生——预算值、实际发生金额、剩余差额；全年累计未发生——预算值、预估发生金额、剩余差额。",
        "再按月份分组（如 12 月、1 月…11 月），每月下列五列：预算值、预估发生、实际发生、完成率、实付。",
        "展示规则：金额与完成率展示规则同前；首条汇总行（如财务指标控制/总能耗汇总）以专用行样式高亮；支持横向滚动；不分页展示。",
        "",
        "界面截图说明",
        "系统实现页面路径：数据概览 → 能耗台账。版式参考 client/src/pages/overview/EnergyLedger.jsx；数据可由能耗台账 Excel 经脚本导出为 energyLedgerData.json 维护。",
    ]

    p = anchor
    for line in lines:
        p = insert_paragraph_after(p, line)

    try:
        doc.save(str(DOC_PATH))
        print(f"已更新: {DOC_PATH}")
    except OSError as e:
        doc.save(str(OUT_PATH))
        print(f"原文件可能被占用，已另存为: {OUT_PATH} ({e})")


if __name__ == "__main__":
    main()
